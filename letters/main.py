import datetime
import random
import docx

class Letter:
    def __init__(self, fullname, workdone, grade, course, pronoun):
        self.fullname = fullname
        self.workdone = workdone
        self.grade = grade
        self.course = course
        self.pronoun = pronoun

    def get_introduction(self):
        x = datetime.datetime.now()
        intro = f"\nStudent: {self.fullname}\n Course: {self.course}\n Date: {x.month}/{x.day}/{x.year}"
        return intro

    def get_course_description(self):
        return "From the start of this course, participants are immersed in the world of data: they are introduced to the concepts of big data, artificial intelligence, the internet of things, cloud computing, and data ethics in the context of real-world business scenarios. Through hands-on experience and practice they study data harvesting and exploration, as well as the basics of data visualization. After they get comfortable with data manipulation and transformation, they gain familiarity with statistical frameworks and methods designed to extract practical insights from data. Participants learn and implement common machine-learning techniques and develop and evaluate analytical solutions."

    def get_content(self):
        firstname = self.fullname.split(' ')[0]
        intro =  f"{self.fullname} took Big Data, Machine Learning and their Real World applications with me this summer.  This was a course for high school students, but it was a college level course."
        details =""
        match self.grade:
            case "A":
                details = self.getRandomAStat(firstname, self.pronoun)
            case "B":
                 details = self.getRandomBStat(firstname, self.pronoun)
            case _:
                details = f"{self.fullname} took Big Data, Machine Learning and their Real World applications with me this summer.  This was a course for high school students, but it was a college level course.  {firstname} started off the course seeming interested in the material and participated in several group projects and discussions throughout the course.  However, {firstname} did not turn in most of yyy assignment prior to the end of the course despite constant reminders to turn in assignments after each presentation. XXX seemed both smart and personable, and, based on yyy efforts in group projects, seemed to have an understanding of the material but due to the lack of homework submissions, I can’t be absolutely sure.  "
        return details + self.workdone

    def getRandomAStat(self, name, pronoun):
        comments = [f"{name} did extremely well in this course. XXX worked hard, participated in class and turned in all of yyy assignments on time.  XXX often went above and beyond the requirements for the assignments, researching and implementing different types of charts than those covered in class. ",
        f"{name} did extremely well in this course and successfully turned in all of yyy assignments on time.  {name} worked very hard and showed a true passion for learning from day one of this course.",
        f"{name} worked very well with yyy teams and put great effort into yyy projects throughout this course. YYY presentations were easy to follow and xxx did a wonderful job using charts and data analytics to find correlations within yyy datasets. {name} had no problem in reaching out to ask questions and gain a deeper understanding of the material. ",
        f"ZZZ presentations were well documented displaying yyy data analysis process as well as the code used to analyze yyy data.  XXX also did an excellent job on all of her assignments. "]
        print(len(comments))
        index = random.randint(0,len(comments)-1)
        return comments[index]

    def getRandomBStat(self, name, pronoun):
        comments =[f"XXX sometimes would wait until an assignment was due to ask more detailed questions about it, which made YYY fall a little behind.  With a little more confidence in asking questions, I do believe that {name} will be able to successfully do college level work. ",
        ]
        index = random.randint(0,len(comments)-1)
        return comments[index]

    def getFinish(self):
        ends =  [f"XXX never hesitated to ask questions when XXX wanted to gain a deeper understanding about machine learning models and how they function.  {self.fullname} deserves an unreserved recommendation to the college of YYY choice. ",
        f" I am confident that xxx would excel at a top-caliber university like Columbia. ",
        f"XXX has the potential to successfully do college-level work.  ",
        f"XXX has thoroughly demonstrated YYY ability to take on and excel at college level work."
        ]
        index = random.randint(0,len(ends)-1)
        return ends[index]
    def pronounify(self,statement):
        pronouned =""
        if self.pronoun == 'she':
            pronouned = statement.replace('xxx','she').replace('yyy', 'her').replace('XXX','She').replace('YYY', 'Her')
        else:
            pronouned = statement.replace('xxx','he').replace('yyy', 'his').replace('XXX','He').replace('YYY', 'His')
        return pronouned


recsToDo = [
    {'name':'Sung Jae Ahn','work':'Sung Jae joined the class already having previous experience with machine learning and neural networks.  He picked up the concepts and syntax of the R programming language very quickly.  For his final project, he tried to incorporate an application that he had built prior to this class in a different language into the language we were using for development in class.  This was very interesting to see as he greatly exceeded the scope of the project but the project presentation would have gone more smoothly if he would have stayed within the predefined scope.', 'grade':'A', 'pronoun':'he'},
    {'name':'Anish Awasthi','work':'Anish did a wonderful job in analyzing and presenting data correlations that he found within multiple datasets.  He also worked with his teammates to build machine learning models that predicted a computer brand based on an image of a computer and the liklihood a student has to succeed in a course based on several factors.', 'grade':'A', 'pronoun':'he'},
    {'name':'Haokai (Daisy) Chen','work':'Haokai worked very hard in this class.  She never hesitated to ask for help or ask questions to gain a deeper understanding of the topics.  She worked very well with her teams and created an app with a machine learning model that predicts an element from the periodic table based on several factors.', 'grade':'A', 'pronoun':'she'},
    {'name':'Tyler Hughes','work':'Tyler did a wonderful job in this course and showed a lot of passion and interest in machine learning.  He worked very well with his teams and created easy to read charts showing the correlations between video game companies and their sales per gaming system. ', 'grade':'A', 'pronoun':'he'},
    {'name':'Isabella Jia','work':'Isabella worked very hard in this course and was eager to ask questions to gain a deeper understanding of the material. She worked well with her team to create a machine learning model that can predict the name of a storm based on several factors. ', 'grade':'A', 'pronoun':'she'},
    {'name':'Aarush Krishnan','work':' He presented all of his findings but surprisingly only turned in a couple of assignments. He worked well with her team to create a machine learning model that can predict the name of a storm based on several factors. I would like to see Aarush be better about turning in his assignments. ', 'grade':'F', 'pronoun':'he'},
    {'name':'Pranav Narayanan','work':'Aarush participated very well in class, always asking questions to gain a deeper understanding of machine learning. He worked well with his team and created a machine learning model application to predict the cost of diamonds based on several factors relating to its quality. ', 'grade':'A', 'pronoun':'he'},
    {'name':'Chuqiao Peng','work':'Chuqiao participated in class very well.  She never hesitated to ask questions and worked well with all other students in class.  Her presentations were very professional, showing easy to read visuals and displaying her correlation findings in an easy to understand way. ', 'grade':'A', 'pronoun':'she'},
    {'name':'Yifeng Shan','work':'Yifeng did very well in this course.  He worked well with his team to do several graphical analyses seeking correlations in dataset and created a machine learning model that can predict the type of cancer based on several factors. ', 'grade':'A', 'pronoun':'he'},
    {'name':'Rayane Souissi','work':'Rayane worked very hard in this course and showed a true passion to learn.  He created a fascinating machine learning model that can predict the level of danger that an object in earths orbit could cause based off of several factors.', 'grade':'A', 'pronoun':'he'},
    {'name':'Yuchen Wei','work':'Yuchen was pretty quiet in class though he did turn in most of his assignments.  ', 'grade':'B', 'pronoun':'he'},   
    {'name':'Lyu Zhuoyang','work':'Lyu was pretty quiet in class but turned in all of his assignment and seemed to have a great understanding of the material. He worked with his team to create a machine learning model application that can predict the type of cancer based on several factors.', 'grade':'A', 'pronoun':'he'},   
    ]

course ='Big Data, Machine Learning, and their Real World Applications'
for rec in recsToDo:
    letter = Letter(rec.get('name'), rec.get('work'), rec.get('grade'), course, rec.get('pronoun'))
    completed = f"{letter.get_introduction()} \n\n {letter.get_course_description()} \n\n {letter.get_content()}  {letter.getFinish()}"
    print(letter.pronounify(completed))
    mydoc =docx.Document()
    mydoc.add_paragraph(letter.pronounify(completed))
    mydoc.add_paragraph("\n\n\nSincerely,")
    mydoc.add_paragraph("Terra Taylor")
    mydoc.add_paragraph(course)
    mydoc.save(f"/Users/terrataylor/Documents/letterwriter/{letter.fullname}.docx")



